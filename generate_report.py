# generate_report.py
# 這個程式會把某一天已經標記好的照片，排成 Word 報告（跟 format.docx 範例檔案一樣的排版）。
#
# 執行方式（在 VS Code 的終端機輸入）：
#     python generate_report.py 2026-04-22
#
# 如果不輸入日期，預設會用「manifest.csv 裡最後一天」的日期。
# 產生的檔案會出現在 output 資料夾裡。

import re
import sys
from datetime import datetime

from docx import Document
from docx.shared import Cm, Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

import utils

RED = RGBColor(0xFF, 0x00, 0x00)

# ---- 排版設定：這些數字是照 format.docx 範例檔量出來的，改動前請先跟範例檔核對 ----
FONT_NAME = "標楷體"

# 頁面邊界（單位：twips，1 公分 = 566.93 twips）
MARGIN_LEFT = Twips(1800)      # 3.175 cm
MARGIN_RIGHT = Twips(1800)     # 3.175 cm
MARGIN_TOP = Twips(1440)       # 2.54 cm
MARGIN_BOTTOM = Twips(1000)    # 約 1.76 cm（原本 2.0 cm；縮小約 0.24 cm 是為了在照片剛好排滿一整頁時，
                                # 讓 Word 規定表格結尾一定要有的收尾段落擠得進同一頁，不然會多出一張
                                # 幾乎全空白、只有頁首的頁）
HEADER_DISTANCE = Twips(851)   # 1.5 cm
FOOTER_DISTANCE = Twips(992)   # 1.75 cm（範例檔沒有頁尾內容，這裡只是保留距離設定）

COLUMN_WIDTH_CM = 8.4
# 照片格子的寬高跟 utils.py 共用同一個常數，app.py 的裁切工具會鎖定一樣的比例
MAX_IMAGE_WIDTH_CM = utils.PHOTO_SLOT_WIDTH_CM
MAX_IMAGE_HEIGHT_CM = utils.PHOTO_SLOT_HEIGHT_CM          # 讓照片留一點邊界，不會頂到列高
IMAGE_ROW_HEIGHT = Twips(3670)     # 範例檔照片列的固定高度，約 6.47 cm
CAPTION_ROW_MIN_HEIGHT = Twips(570)  # 範例檔說明文字列的高度，約 1.0 cm
PHOTOS_PER_TABLE = 6                # 每個分類固定 3 對（6 張），不夠就留空白格


def roc_date(date_str: str) -> str:
    """把 2026-04-22 轉成民國日期 115年04月22日"""
    d = datetime.strptime(date_str, "%Y-%m-%d")
    roc_year = d.year - 1911
    return f"{roc_year}年{d.month:02d}月{d.day:02d}日"


def set_font(run, name=FONT_NAME, size=12, bold=False, italic=False, color=None):
    """套用範例檔用的字型（標楷體），並確保中文（東亞字元）也套用到同一個字型。"""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def set_table_cell_margins(table, top=0, left=28, bottom=0, right=28):
    """套用範例檔裡表格儲存格的內距（左右幾乎貼齊，跟預設的 Table Grid 樣式不同）。"""
    tblPr = table._tbl.tblPr
    tblCellMar = OxmlElement("w:tblCellMar")
    for side, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tblCellMar.append(node)
    tblPr.append(tblCellMar)


def set_row_height(row, height, rule):
    row.height = height
    row.height_rule = rule


def add_fitted_picture(run, photo_path):
    """把照片縮放到剛好放進固定的照片列高度內，維持長寬比，不會被裁切。"""
    with Image.open(photo_path) as img:
        w_px, h_px = img.size
    aspect = w_px / h_px
    width_cm = MAX_IMAGE_WIDTH_CM
    height_cm = width_cm / aspect
    if height_cm > MAX_IMAGE_HEIGHT_CM:
        height_cm = MAX_IMAGE_HEIGHT_CM
        width_cm = height_cm * aspect
    run.add_picture(str(photo_path), width=Cm(width_cm), height=Cm(height_cm))


def add_segmented_text(paragraph, text, red_terms=None, size=12, bold=False):
    """把 text 依 red_terms（例如鋪築施工照片的材料代號 DGAC/OGAC/PAC）切成好幾個 run，
    命中的片段用紅字，其餘維持正常黑字。回傳最後一個 run，方便呼叫端接著 add_break 加註記。"""
    if not red_terms:
        run = paragraph.add_run(text)
        set_font(run, size=size, bold=bold)
        return run

    pattern = "(" + "|".join(re.escape(t) for t in red_terms) + ")"
    last_run = None
    for part in re.split(pattern, text):
        if not part:
            continue
        run = paragraph.add_run(part)
        set_font(run, size=size, bold=bold, color=RED if part in red_terms else None)
        last_run = run
    if last_run is None:
        last_run = paragraph.add_run("")
        set_font(last_run, size=size, bold=bold)
    return last_run


def set_cell_text(cell, text, bold=False, size=12, align_center=True, vertical_center=True, note=None,
                   red_terms=None, note_color=None):
    cell.paragraphs[0].text = ""
    p = cell.paragraphs[0]
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if vertical_center:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    last_run = add_segmented_text(p, text, red_terms=red_terms, size=size, bold=bold)
    if note:
        # 換行加註一行較小、斜體的說明文字（例如：拍照日期跟施工日期不同時，或分類張數不足的警示）
        last_run.add_break(WD_BREAK.LINE)
        note_run = p.add_run(note)
        set_font(note_run, size=max(8, size - 2), italic=True, color=note_color)


def build_report(date_str: str, config: dict, rows: list):
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.header_distance = HEADER_DISTANCE
    section.footer_distance = FOOTER_DISTANCE

    # ---- 頁首（每一頁都會重複出現，跟 format.docx 一樣） ----
    header = section.header
    header.is_linked_to_previous = False

    title_p = header.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.line_spacing = 1.0
    set_font(title_p.add_run(config["project_name"]), size=16)

    subtitle_p = header.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.line_spacing = 1.0
    set_font(subtitle_p.add_run(config["report_title"]), size=16)

    info_p = header.add_paragraph()
    info_p.paragraph_format.line_spacing = 1.0
    info_p.paragraph_format.space_after = Pt(0)
    # 用「靠右對齊的定位點」取代手動打空格：不管承攬廠商名稱多長，
    # 施工日期都會貼齊右邊界，不會因為名稱長度改變而跑位。
    usable_width = section.page_width - section.left_margin - section.right_margin
    info_p.paragraph_format.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)
    set_font(info_p.add_run(f"承攬廠商：{config['contractor']}"), size=12)
    info_p.add_run("\t")
    set_font(info_p.add_run(f"施工日期：{roc_date(date_str)}"), size=12)

    # ---- 依分類分組 ----
    rows_by_category = {}
    for row in rows:
        rows_by_category.setdefault(row["category"], []).append(row)

    any_photos = False
    is_first_table = True
    for cat in config["categories"]:
        cat_name = cat["name"]
        cat_display_name = utils.display_name(cat)
        cat_rows = rows_by_category.get(cat_name, [])
        if not cat_rows:
            continue
        any_photos = True

        # 同一分類底下，照片要依照 config.json 設定的說明順序排列，不是依照標記的先後順序。
        # 找不到對應說明（例如自行輸入的自訂說明）就放到最後面，彼此之間維持原本標記的先後順序。
        def sort_key(row):
            slot = utils.caption_slot_index(row["caption"], cat)
            return (slot is None, slot or 0)

        cat_rows = sorted(cat_rows, key=sort_key)

        material_options = cat.get("material_options")
        min_required = cat.get("min_required")
        shortage_note = None
        if min_required and len(cat_rows) < min_required:
            shortage_note = f"⚠️ 僅 {len(cat_rows)} 張，未達最低 {min_required} 張要求"
            print(f"警告：{cat_name} 在 {date_str} 只有 {len(cat_rows)} 張，未達最低 {min_required} 張要求！")

        # 每個分類固定 3 對（6 張）一頁：不夠就補空白格，超過就再開一頁，
        # 换分类或换页都强制分頁，不會跟其他分類擠在同一頁。
        for chunk_start in range(0, len(cat_rows), PHOTOS_PER_TABLE):
            chunk = cat_rows[chunk_start:chunk_start + PHOTOS_PER_TABLE]
            chunk = chunk + [None] * (PHOTOS_PER_TABLE - len(chunk))

            table = doc.add_table(rows=0, cols=2)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            set_table_cell_margins(table)

            # 分類標題（合併整列）：跟說明文字列用一樣的最小高度，垂直空間才會一致
            header_row = table.add_row()
            header_row.cells[0].merge(header_row.cells[1])
            set_row_height(header_row, CAPTION_ROW_MIN_HEIGHT, WD_ROW_HEIGHT_RULE.AT_LEAST)
            set_cell_text(header_row.cells[0], cat_display_name, bold=False, size=12, note=shortage_note, note_color=RED)
            if not is_first_table:
                # 用「段落屬性：之前分頁」讓這個表格自己另起一頁，而不是在前一個表格後面另外加一個
                # 換頁段落。如果前一個表格剛好整頁塞滿，額外加的那個換頁段落會被擠到下一頁才顯示，
                # 段落裡的分頁字元又再往後推一頁，等於多出一整張空白頁。
                header_row.cells[0].paragraphs[0].paragraph_format.page_break_before = True
            is_first_table = False

            for i in range(0, PHOTOS_PER_TABLE, 2):
                pair = chunk[i:i + 2]

                img_row = table.add_row()
                cap_row = table.add_row()
                set_row_height(img_row, IMAGE_ROW_HEIGHT, WD_ROW_HEIGHT_RULE.EXACTLY)
                set_row_height(cap_row, CAPTION_ROW_MIN_HEIGHT, WD_ROW_HEIGHT_RULE.AT_LEAST)

                for col_idx in range(2):
                    img_row.cells[col_idx].width = Cm(COLUMN_WIDTH_CM)
                    cap_row.cells[col_idx].width = Cm(COLUMN_WIDTH_CM)

                    item = pair[col_idx]
                    if item is not None:
                        # manifest 裡存的是「相對於這個工地資料夾」的路徑（見 app.py 寫入處）
                        photo_path = utils.DATA_DIR / item["sorted_path"]

                        img_cell = img_row.cells[col_idx]
                        img_cell.paragraphs[0].text = ""
                        p = img_cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        if photo_path.exists():
                            add_fitted_picture(run, photo_path)
                        else:
                            run.text = f"[找不到照片: {photo_path.name}]"
                            set_font(run, size=11)

                        set_cell_text(
                            cap_row.cells[col_idx], item["caption"], size=12,
                            red_terms=material_options,
                        )
                    else:
                        set_cell_text(img_row.cells[col_idx], "")
                        set_cell_text(cap_row.cells[col_idx], "")

    if not any_photos:
        doc.add_paragraph("（這個日期目前沒有已標記的照片）")
    else:
        # Word 規定表格不能是文件最後一個元素，一定要有一個段落收尾，不然 Word 自己補一個
        # 預設樣式的段落時，行高/間距比這裡故意留小，常常會被擠到最後一頁後面多一張幾乎空白的頁。
        # 這裡自己加一個字級跟間距都設到最小的段落，盡量讓它留在最後一個表格的同一頁。
        trailing_p = doc.add_paragraph()
        trailing_p.paragraph_format.space_before = Pt(0)
        trailing_p.paragraph_format.space_after = Pt(0)
        trailing_p.paragraph_format.line_spacing = 1.0
        set_font(trailing_p.add_run(""), size=1)

    utils.OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    # 有設定工地名稱時（同一天有兩個以上工地在施工）就把名稱寫進檔名。兩個工地各自有自己的
    # output 資料夾，本來就不會互相覆蓋，但報告產生後通常會被複製到同一個資料夾或寄出去，
    # 那時候兩個「查驗照片_2026-07-23.docx」就會直接蓋掉對方，而且看不出蓋掉的是哪一個。
    # 只有一個工地（沒設定名稱）時維持原本的檔名不變。
    site_suffix = f"_{utils.safe_filename(utils.SITE_NAME)}" if utils.SITE_NAME else ""
    out_path = utils.OUTPUT_DIR / f"{utils.REPORT_FILENAME_PREFIX}{date_str}{site_suffix}.docx"
    doc.save(out_path)
    return out_path


def main():
    utils.ensure_dirs()
    config = utils.load_config()
    all_rows = utils.load_manifest()

    if not all_rows:
        print("manifest.csv 裡還沒有任何已標記的照片。請先執行 app.py 標記照片。")
        return

    if len(sys.argv) > 1:
        date_str = sys.argv[1]
    else:
        date_str = sorted({r["date"] for r in all_rows})[-1]
        print(f"未指定日期，使用最新日期：{date_str}")

    day_rows = [r for r in all_rows if r["date"] == date_str]
    if not day_rows:
        print(f"找不到日期 {date_str} 的任何已標記照片。")
        return

    out_path = build_report(date_str, config, day_rows)
    print(f"完成！報告已產生：{out_path}")


if __name__ == "__main__":
    main()
